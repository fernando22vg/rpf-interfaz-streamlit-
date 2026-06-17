"""
reporte_tecnico.py
------------------
Generador de Reporte Técnico RPF para el Bloque 7 de la interfaz.

Funciones principales:
  · collect_kpis_para_reporte() — recolecta KPIs de todas las fuentes
  · generar_reporte_word()      — genera documento .docx con python-docx
  · generar_copilot_context_pack() — texto estructurado + prompts para Copilot
  · subir_reporte_a_sharepoint() — sube el .docx al SharePoint corporativo

IMPORTANTE: Este módulo NO importa Streamlit. Es seguro para uso fuera del
contexto de la app Streamlit (excepto _cargar_historico_unidad que usa
bloque_kpi_historico y llama a st.cache_data cuando corre dentro de Streamlit).
"""

from __future__ import annotations

import glob as _glob
import io
import os
import tempfile
from dataclasses import dataclass, field
from datetime import datetime
from typing import Optional

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from kpi_calc import (
    _calcular_rocof,
    _cndc_kpis,
    _get_pmax_from_cargado,
    _get_rp_default,
    _load_pmax_cargado,
    _load_tech_map,
    _robust_col_detect,
)

# ── Constantes ────────────────────────────────────────────────────────────────

_CARPETA_SCADA   = "Graficas Registro 1SEG COBEE"
_CARPETA_EMF     = "Resultados_COBEE"
_CARPETA_SIM     = "Datos Curvas"

_COL_COBEE_AZUL  = "2E5C8A"
_COL_OK_BG       = "D1FAE5"  # verde claro
_COL_NOK_BG      = "FEE2E2"  # rojo claro
_COL_HEADER_FG   = RGBColor(0xFF, 0xFF, 0xFF)
_COL_PLACEHOLDER = RGBColor(0x94, 0xA3, 0xB8)  # gris para texto Copilot

_KPI_LABELS = [
    ("P_max [MW]",                  "p_max_val"),
    ("f₀ — Inicio evento [Hz]",     "f0"),
    ("P₀ — Inicio evento [MW]",     "p0"),
    ("f_min — Nadir [Hz]",          "f_min"),
    ("t_min — Tiempo nadir [s]",    "t_min"),
    ("Δf = f₀ − f_min [Hz]",       "delta_f"),
    ("f_Δt [Hz]",                   "f_dt"),
    ("P_Δt [MW]",                   "p_dt"),
    ("Reserva Inicial R_inic [MW]", "r_inic"),
    ("Reserva Inicial R_inic [%]",  "r_inic_pct"),
    ("ΔP entregada [MW]",           "dp"),
    ("ΔP% aporte [%]",              "dp_pct"),
    ("¿Aporta a la RPF?",           "aporta"),
    ("Droop Nominal [%]",           "droop_nom"),
    ("Droop Calculado [%]",         "droop_calc"),
    ("ROCOF [Hz/s]",                "rocof"),
]


# ── Dataclass central ─────────────────────────────────────────────────────────

@dataclass
class ReporteContext:
    semestre:    str
    n_evento:    int | str
    unit_name:   str        # nombre limpio sin "sym_"
    unit_key:    str        # loc_name PF original
    p_max:       float      # MW
    rp_nominal:  float      # fracción (0.05 = 5 %)
    delta_t:     int        # segundos
    t_falla:     float      # segundos en eje de simulación
    ev_path:     str

    kpi_scada:   Optional[dict] = None
    kpi_emf:     Optional[dict] = None
    kpi_sim_e0:  Optional[dict] = None
    kpi_sim_e1:  Optional[dict] = None

    historico:   pd.DataFrame = field(default_factory=pd.DataFrame)
    audit:       dict = field(default_factory=dict)


# ── Helpers duplicados (sin importar interfaz_analisis_RPF.py) ─────────────────

def _parse_to_seconds(series: pd.Series) -> pd.Series:
    s = series.astype(str).str.strip().str.replace(",", ".", regex=False)
    result = pd.Series(np.nan, index=series.index)
    has_colon = s.str.contains(":")
    if has_colon.any():
        parts = s[has_colon].str.split(":")
        h   = pd.to_numeric(parts.str[0], errors="coerce").fillna(0)
        m   = pd.to_numeric(parts.str[1], errors="coerce").fillna(0)
        sec = pd.to_numeric(parts.str[2], errors="coerce").fillna(0)
        result[has_colon] = h * 3600 + m * 60 + sec
    not_colon = ~has_colon
    if not_colon.any():
        result[not_colon] = pd.to_numeric(s[not_colon], errors="coerce")
    return result.fillna(0.0)


def _detectar_inicio_falla(freq_array: np.ndarray,
                            umbral_dfdt: float = -0.02,
                            ventana_suavizado: int = 5) -> int:
    n = len(freq_array)
    if n < ventana_suavizado + 2:
        return 0
    kernel = np.ones(ventana_suavizado) / ventana_suavizado
    freq_smooth = np.convolve(freq_array.astype(float), kernel, mode="same")
    freq_smooth[:ventana_suavizado // 2] = freq_smooth[ventana_suavizado // 2]
    freq_smooth[-(ventana_suavizado // 2):] = freq_smooth[-(ventana_suavizado // 2) - 1]
    dfdt = np.diff(freq_smooth)
    condicion = (dfdt[:-1] < umbral_dfdt) & (dfdt[1:] < umbral_dfdt)
    indices = np.where(condicion)[0]
    if len(indices) > 0:
        return int(indices[0] + 1)
    candidatos = np.where(dfdt < umbral_dfdt)[0]
    return int(candidatos[0] + 1) if len(candidatos) > 0 else 0


def _buscar_archivo_unidad(unit_name: str, file_list: list) -> Optional[str]:
    if not unit_name:
        return None
    u_norm = str(unit_name).upper().replace("SYM_", "").replace("SYM", "")
    for f in file_list:
        f_base = os.path.splitext(f)[0].upper().replace("SYM_", "").replace("SYM", "")
        if u_norm == f_base:
            return f
    for f in file_list:
        f_base = os.path.splitext(f)[0].upper().replace("SYM_", "").replace("SYM", "")
        if u_norm in f_base or f_base in u_norm:
            return f
    return None


# ── Lectura de arrays (sin @st.cache_data) ────────────────────────────────────

def _read_sim_arrays(file_path: str, t_falla: float):
    """Carga arrays de simulación PowerFactory. Devuelve (ts_aligned, fs_hz, ps_mw)."""
    df = pd.read_excel(file_path, engine="calamine").dropna()
    tc, fc, pc = _robust_col_detect(df)
    ts_raw = pd.to_numeric(df[tc], errors="coerce").values
    fs_raw = pd.to_numeric(df[fc], errors="coerce").ffill().values
    fs_hz  = fs_raw * 50.0 if np.nanmax(fs_raw) < 2.0 else fs_raw
    ps_mw  = pd.to_numeric(df[pc], errors="coerce").ffill().values
    valid  = ~np.isnan(ts_raw)
    return ts_raw[valid] - t_falla, fs_hz[valid], ps_mw[valid]


def _read_real_arrays(file_path: str,
                      umbral_dfdt: float = 0.05,
                      ventana: int = 5):
    """Carga arrays de datos reales (SCADA/EMF). Devuelve (ts_aligned, fs_hz, ps_mw)."""
    df     = pd.read_excel(file_path, engine="calamine").dropna()
    tr_raw = _parse_to_seconds(df.iloc[:, 0])
    tr_norm = tr_raw - tr_raw.min()
    _fr_c   = [c for c in df.columns if any(kw in c.lower() for kw in ["frec", "hz", "freq"])]
    _fr_col = _fr_c[0] if _fr_c else df.columns[1]
    _pr_col = df.columns[2] if len(df.columns) > 2 else df.columns[1]
    fr_arr  = pd.to_numeric(df[_fr_col], errors="coerce").ffill().values
    pr_arr  = pd.to_numeric(df[_pr_col], errors="coerce").ffill().values
    idx_f   = _detectar_inicio_falla(fr_arr, umbral_dfdt, ventana)
    t_f     = float(tr_norm.iloc[idx_f])
    return (tr_norm - t_f).values, fr_arr, pr_arr


# ── KPIs por fuente ───────────────────────────────────────────────────────────

def _kpi_from_sim(ev_path: str, n_evento, unit_key: str,
                  sim_ver: str, t_falla: float,
                  p_max: float, rp: float, delta_t: int,
                  loc_gen_path: str) -> Optional[dict]:
    sim_dir = os.path.join(ev_path, f"E{n_evento}.{sim_ver}", _CARPETA_SIM)
    if not os.path.isdir(sim_dir):
        return None
    xlsx_files = [f for f in os.listdir(sim_dir) if f.endswith(".xlsx") and not f.startswith("~$")]
    match = _buscar_archivo_unidad(unit_key, xlsx_files)
    if not match:
        return None
    fpath = os.path.join(sim_dir, match)
    try:
        ts, fs, ps = _read_sim_arrays(fpath, t_falla)
        kpi = _cndc_kpis(ts, fs, ps, p_max, rp, delta_t)
        if kpi is None:
            return None
        kpi["rocof"]  = _calcular_rocof(ts, fs)
        kpi["fuente"] = f"Simulación E{n_evento}.{sim_ver}"
        kpi["p_max_val"] = p_max
        return kpi
    except Exception:
        return None


def _kpi_from_real(filepath: str, p_max: float, rp: float,
                   delta_t: int, umbral_dfdt: float, ventana: int,
                   fuente_nombre: str) -> Optional[dict]:
    if not os.path.isfile(filepath):
        return None
    try:
        ts, fs, ps = _read_real_arrays(filepath, umbral_dfdt, ventana)
        kpi = _cndc_kpis(ts, fs, ps, p_max, rp, delta_t)
        if kpi is None:
            return None
        kpi["rocof"]  = _calcular_rocof(ts, fs)
        kpi["fuente"] = fuente_nombre
        kpi["p_max_val"] = p_max
        return kpi
    except Exception:
        return None


def _cargar_historico_unidad(unit_name: str) -> pd.DataFrame:
    """Carga el histórico RPF de la unidad desde bloque_kpi_historico._load_data()."""
    try:
        from bloque_kpi_historico import _load_data  # solo disponible en contexto Streamlit
        df, _ = _load_data()
        if df.empty:
            return pd.DataFrame()
        u = unit_name.upper().replace("SYM_", "").replace("SYM", "")
        mask = df["unidad"].astype(str).str.upper().str.replace("SYM_", "").str.replace("SYM", "")
        return df[mask == u].copy()
    except Exception:
        return pd.DataFrame()


# ── Función principal de recolección ─────────────────────────────────────────

def collect_kpis_para_reporte(
    ev_path:      str,
    n_evento,
    unit_key:     str,
    delta_t:      int,
    t_falla:      float,
    loc_gen_path: str,
    umbral_dfdt:  float = 0.05,
    ventana_det:  int   = 5,
) -> ReporteContext:
    """
    Recolecta KPIs de todas las fuentes disponibles para la unidad.
    Cada fuente se intenta de forma independiente; si falla, kpi_xxx = None.
    """
    unit_name = os.path.splitext(unit_key)[0].replace("sym_", "").replace("SYM_", "")

    # P_max y Rp
    try:
        pmax_cargado = _load_pmax_cargado(ev_path, n_evento)
        tech_map     = _load_tech_map(loc_gen_path)
        p_max, tk, _ = _get_pmax_from_cargado(unit_key, pmax_cargado, tech_map)
        rp = float(_get_rp_default(tk, loc_gen_path)) / 100.0
    except Exception:
        p_max, rp, tk = 200.0, 0.05, unit_key

    # Auditoría de fuentes
    scada_dir = os.path.join(ev_path, _CARPETA_SCADA)
    emf_dir   = os.path.join(ev_path, _CARPETA_EMF)
    audit = {
        "SCADA disponible": os.path.isdir(scada_dir),
        "EMF disponible":   os.path.isdir(emf_dir),
        f"Sim E{n_evento}.0 disponible": os.path.isdir(os.path.join(ev_path, f"E{n_evento}.0", _CARPETA_SIM)),
        f"Sim E{n_evento}.1 disponible": os.path.isdir(os.path.join(ev_path, f"E{n_evento}.1", _CARPETA_SIM)),
    }

    # KPIs SCADA
    kpi_scada = None
    if os.path.isdir(scada_dir):
        scada_files = [f for f in os.listdir(scada_dir) if f.endswith(".xlsx") and not f.startswith("~$")]
        match_s = _buscar_archivo_unidad(unit_key, scada_files)
        if match_s:
            kpi_scada = _kpi_from_real(
                os.path.join(scada_dir, match_s), p_max, rp, delta_t,
                umbral_dfdt, ventana_det, "SCADA COBEE (1 SEG)"
            )

    # KPIs EMF
    kpi_emf = None
    if os.path.isdir(emf_dir):
        emf_files = [f for f in os.listdir(emf_dir) if f.endswith(".xlsx") and not f.startswith("~$")]
        match_e = _buscar_archivo_unidad(unit_key, emf_files)
        if match_e:
            kpi_emf = _kpi_from_real(
                os.path.join(emf_dir, match_e), p_max, rp, delta_t,
                umbral_dfdt, ventana_det, "EMF CNDC (Digitalización)"
            )

    # KPIs Simulación E.0
    kpi_sim_e0 = _kpi_from_sim(ev_path, n_evento, unit_key, "0", t_falla, p_max, rp, delta_t, loc_gen_path)

    # KPIs Simulación E.1
    kpi_sim_e1 = _kpi_from_sim(ev_path, n_evento, unit_key, "1", t_falla, p_max, rp, delta_t, loc_gen_path)

    # Histórico
    historico = _cargar_historico_unidad(unit_name)

    return ReporteContext(
        semestre   = "",   # se rellena en el Bloque 7
        n_evento   = n_evento,
        unit_name  = unit_name,
        unit_key   = unit_key,
        p_max      = float(p_max),
        rp_nominal = rp,
        delta_t    = delta_t,
        t_falla    = t_falla,
        ev_path    = ev_path,
        kpi_scada  = kpi_scada,
        kpi_emf    = kpi_emf,
        kpi_sim_e0 = kpi_sim_e0,
        kpi_sim_e1 = kpi_sim_e1,
        historico  = historico,
        audit      = audit,
    )


# ── Helpers Word ──────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    """Aplica color de fondo a una celda Word vía XML."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _cell_text(cell, text: str, bold: bool = False,
               size_pt: int = 10, color: RGBColor = None,
               italic: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    """Escribe texto en una celda con formato."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color


def _fmt_kpi_value(field_name: str, kpi: dict) -> str:
    """Formatea el valor de un campo KPI para mostrar en tabla."""
    val = kpi.get(field_name)
    if field_name == "p_max_val":
        return f"{float(val):.2f}" if val is not None else "—"
    if field_name == "aporta":
        return "✔ Sí" if val else "✘ No"
    if field_name == "rocof":
        if val is None or (isinstance(val, float) and np.isnan(val)):
            return "—"
        return f"{float(val):.4f}"
    if val is None:
        return "—"
    if isinstance(val, float) and np.isnan(val):
        return "—"
    if isinstance(val, float):
        return f"{val:.4f}" if abs(val) < 100 and "pct" not in field_name and field_name not in ("t_min", "droop_nom", "droop_calc") else f"{val:.2f}"
    return str(val)


def _add_header_row(table, headers: list[str]) -> None:
    """Agrega fila de encabezado con fondo azul COBEE y texto blanco."""
    row = table.rows[0]
    for i, h in enumerate(headers):
        c = row.cells[i]
        _set_cell_bg(c, _COL_COBEE_AZUL)
        _cell_text(c, h, bold=True, size_pt=9, color=_COL_HEADER_FG,
                   align=WD_ALIGN_PARAGRAPH.CENTER)


def _add_placeholder_copilot(doc: Document, instruccion: str) -> None:
    """Agrega párrafo gris-cursiva como guía para Copilot."""
    p = doc.add_paragraph()
    run = p.add_run(f"[COPILOT] {instruccion}")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = _COL_PLACEHOLDER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)


# ── Secciones del documento Word ──────────────────────────────────────────────

def _doc_portada(doc: Document, ctx: ReporteContext) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("COBEE S.A.")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)

    doc.add_paragraph()

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("INFORME DE RESERVA DE POTENCIA DE FRECUENCIA (RPF)")
    run2.bold = True
    run2.font.size = Pt(16)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(f"Sistema Interconectado Nacional de Bolivia")
    run3.font.size = Pt(11)

    doc.add_paragraph()

    # Tabla de identificación
    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = "Table Grid"
    datos = [
        ("Semestre:", ctx.semestre or "—"),
        ("Evento:",   f"Evento {ctx.n_evento}"),
        ("Unidad Generadora:", ctx.unit_name),
        ("Fecha del informe:", datetime.now().strftime("%d/%m/%Y %H:%M")),
        ("Elaborado por:", "COBEE S.A. — Dpto. Ingeniería"),
    ]
    for i, (label, valor) in enumerate(datos):
        _set_cell_bg(tbl.cell(i, 0), _COL_COBEE_AZUL)
        _cell_text(tbl.cell(i, 0), label, bold=True, size_pt=10, color=_COL_HEADER_FG)
        _cell_text(tbl.cell(i, 1), valor, size_pt=10)

    doc.add_page_break()


def _doc_resumen_ejecutivo(doc: Document, ctx: ReporteContext) -> None:
    doc.add_heading("1. Resumen Ejecutivo", level=1)

    fuentes = [
        ("SCADA COBEE (1 SEG)",     ctx.kpi_scada),
        ("EMF CNDC",                ctx.kpi_emf),
        (f"Sim E{ctx.n_evento}.0",  ctx.kpi_sim_e0),
        (f"Sim E{ctx.n_evento}.1",  ctx.kpi_sim_e1),
    ]
    disponibles = [(n, k) for n, k in fuentes if k is not None]

    # Tabla semáforo
    headers = ["Fuente de Datos", "ΔP%", "Droop Calculado [%]", "¿Aporta RPF?"]
    tbl = doc.add_table(rows=1 + len(fuentes), cols=4)
    tbl.style = "Table Grid"
    _add_header_row(tbl, headers)
    for i, (nombre, kpi) in enumerate(fuentes, start=1):
        row = tbl.rows[i]
        if kpi is None:
            _cell_text(row.cells[0], nombre, size_pt=9)
            for c in row.cells[1:]:
                _cell_text(c, "No disponible", italic=True, size_pt=9,
                           color=_COL_PLACEHOLDER, align=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            aporta = kpi.get("aporta", False)
            bg = _COL_OK_BG if aporta else _COL_NOK_BG
            _cell_text(row.cells[0], nombre, size_pt=9)
            _cell_text(row.cells[1], f"{kpi.get('dp_pct', 0):.2f}%",
                       size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            droop_c = kpi.get("droop_calc", "—")
            _cell_text(row.cells[2], f"{droop_c:.1f}%" if isinstance(droop_c, float) else str(droop_c),
                       size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell_text(row.cells[3], "✔ Sí" if aporta else "✘ No",
                       bold=True, size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            for c in row.cells[1:]:
                _set_cell_bg(c, bg)

    doc.add_paragraph()
    _add_placeholder_copilot(
        doc,
        "Redacta un párrafo ejecutivo (máx. 100 palabras) resumiendo si la unidad cumplió "
        "con la RPF, el nadir alcanzado y el droop calculado según los datos de la tabla anterior."
    )


def _doc_contexto_evento(doc: Document, ctx: ReporteContext) -> None:
    doc.add_heading("2. Contexto del Evento", level=1)
    doc.add_heading("2.1 Identificación de la Unidad", level=2)

    tbl = doc.add_table(rows=6, cols=2)
    tbl.style = "Table Grid"
    filas = [
        ("Semestre de Análisis",    ctx.semestre or "—"),
        ("Número de Evento",        f"Evento {ctx.n_evento}"),
        ("Unidad Generadora",       ctx.unit_name),
        ("Potencia Efectiva P_max", f"{ctx.p_max:.2f} MW"),
        ("Estatismo Nominal (Rp)",  f"{ctx.rp_nominal * 100:.1f}%"),
        ("Ventana de Evaluación Δt", f"{ctx.delta_t} s"),
    ]
    for i, (label, val) in enumerate(filas):
        _set_cell_bg(tbl.cell(i, 0), "EFF6FF")
        _cell_text(tbl.cell(i, 0), label, bold=True, size_pt=9)
        _cell_text(tbl.cell(i, 1), val, size_pt=9)

    doc.add_heading("2.2 Auditoría de Fuentes Disponibles", level=2)
    audit_rows = [
        ("SCADA COBEE (1 SEG)",     ctx.kpi_scada is not None),
        ("EMF CNDC (Digitalización)", ctx.kpi_emf is not None),
        (f"Simulación E{ctx.n_evento}.0 (CNDC)", ctx.kpi_sim_e0 is not None),
        (f"Simulación E{ctx.n_evento}.1 (COBEE)", ctx.kpi_sim_e1 is not None),
    ]
    tbl2 = doc.add_table(rows=1 + len(audit_rows), cols=3)
    tbl2.style = "Table Grid"
    _add_header_row(tbl2, ["Fuente de Datos", "Estado", "Observación"])
    for i, (nombre, ok) in enumerate(audit_rows, start=1):
        bg = _COL_OK_BG if ok else _COL_NOK_BG
        _cell_text(tbl2.cell(i, 0), nombre, size_pt=9)
        _set_cell_bg(tbl2.cell(i, 1), bg)
        _cell_text(tbl2.cell(i, 1), "Incluida" if ok else "No disponible",
                   bold=True, size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(tbl2.cell(i, 2),
                   "Archivo encontrado y procesado" if ok else "Archivo no encontrado para esta unidad",
                   italic=True, size_pt=9)


def _doc_tabla_kpi_fuente(doc: Document, fuente_nombre: str,
                           kpi: Optional[dict], p_max: float, delta_t: int) -> None:
    """Inserta una subsección con tabla KPI completa para una fuente."""
    if kpi is None:
        p = doc.add_paragraph(f"⚠ Fuente no disponible para esta unidad en este evento.")
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = _COL_PLACEHOLDER
        return

    tbl = doc.add_table(rows=1 + len(_KPI_LABELS), cols=2)
    tbl.style = "Table Grid"
    _add_header_row(tbl, ["Indicador CNDC", f"Valor — {fuente_nombre}"])

    for i, (label, field_key) in enumerate(_KPI_LABELS, start=1):
        row = tbl.rows[i]
        _cell_text(row.cells[0], label, size_pt=9)
        val_str = _fmt_kpi_value(field_key, kpi)
        # Fondo verde/rojo en la fila ΔP% y ¿Aporta?
        if field_key in ("dp_pct", "aporta"):
            aporta = kpi.get("aporta", False)
            _set_cell_bg(row.cells[0], _COL_OK_BG if aporta else _COL_NOK_BG)
            _set_cell_bg(row.cells[1], _COL_OK_BG if aporta else _COL_NOK_BG)
            _cell_text(row.cells[1], val_str, bold=True, size_pt=9,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            _cell_text(row.cells[1], val_str, size_pt=9,
                       align=WD_ALIGN_PARAGRAPH.CENTER)


def _doc_resultados_por_fuente(doc: Document, ctx: ReporteContext) -> None:
    doc.add_heading("3. Resultados por Fuente de Datos", level=1)

    fuentes = [
        ("3.1 SCADA COBEE (1 SEG)",              ctx.kpi_scada),
        ("3.2 EMF CNDC (Digitalización)",         ctx.kpi_emf),
        (f"3.3 Simulación E{ctx.n_evento}.0 — CNDC", ctx.kpi_sim_e0),
        (f"3.4 Simulación E{ctx.n_evento}.1 — COBEE", ctx.kpi_sim_e1),
    ]
    for heading, kpi in fuentes:
        doc.add_heading(heading, level=2)
        fuente_corta = heading.split(" ", 1)[1] if " " in heading else heading
        _doc_tabla_kpi_fuente(doc, fuente_corta, kpi, ctx.p_max, ctx.delta_t)
        doc.add_paragraph()


def _doc_tabla_comparativa(doc: Document, ctx: ReporteContext) -> None:
    doc.add_heading("4. Análisis Comparativo entre Fuentes", level=1)

    nombres = [
        "SCADA", "EMF",
        f"Sim E{ctx.n_evento}.0", f"Sim E{ctx.n_evento}.1"
    ]
    kpis = [ctx.kpi_scada, ctx.kpi_emf, ctx.kpi_sim_e0, ctx.kpi_sim_e1]

    # Solo mostrar columnas de fuentes con datos
    activos = [(n, k) for n, k in zip(nombres, kpis) if k is not None]
    if not activos:
        doc.add_paragraph("No hay fuentes de datos disponibles para comparar.")
        return

    campos_comp = [
        ("f₀ [Hz]",        "f0"),
        ("P₀ [MW]",        "p0"),
        ("f_min [Hz]",     "f_min"),
        ("t_min [s]",      "t_min"),
        ("Δf [Hz]",        "delta_f"),
        ("ΔP [MW]",        "dp"),
        ("ΔP% [%]",        "dp_pct"),
        ("Droop calc [%]", "droop_calc"),
        ("ROCOF [Hz/s]",   "rocof"),
    ]

    ncols = 1 + len(activos)
    tbl = doc.add_table(rows=1 + len(campos_comp), cols=ncols)
    tbl.style = "Table Grid"
    _add_header_row(tbl, ["Indicador"] + [n for n, _ in activos])

    for i, (label, fkey) in enumerate(campos_comp, start=1):
        row = tbl.rows[i]
        _cell_text(row.cells[0], label, bold=True, size_pt=9)
        for j, (_, kpi) in enumerate(activos, start=1):
            val = _fmt_kpi_value(fkey, kpi)
            if fkey == "dp_pct":
                aporta = kpi.get("aporta", False)
                _set_cell_bg(row.cells[j], _COL_OK_BG if aporta else _COL_NOK_BG)
            _cell_text(row.cells[j], val, size_pt=9,
                       align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    _add_placeholder_copilot(
        doc,
        "Analiza las diferencias entre las fuentes disponibles. "
        "¿Son consistentes los resultados? ¿Qué factores pueden explicar las discrepancias "
        "entre SCADA/EMF y los modelos de simulación en ΔP% y droop calculado?"
    )


def _doc_historico(doc: Document, ctx: ReporteContext) -> None:
    doc.add_heading("5. Contexto Histórico de la Unidad", level=1)

    hist = ctx.historico
    if hist.empty:
        doc.add_paragraph("No se encontraron datos históricos para esta unidad en la base de datos.")
        _add_placeholder_copilot(
            doc,
            "Con base en los resultados del evento actual, ¿qué tendencia histórica esperarías "
            "para esta unidad? ¿Los valores de ΔP% y droop son coherentes con un generador de su tipo?"
        )
        return

    cols_show = ["semestre", "evento", "f0", "f_min", "dp_pct", "droop_calc", "aporta"]
    cols_disp = [c for c in cols_show if c in hist.columns]
    hist_sub   = hist[cols_disp].tail(10)

    headers_map = {
        "semestre": "Semestre", "evento": "Evento",
        "f0": "f₀ [Hz]", "f_min": "f_min [Hz]",
        "dp_pct": "ΔP% [%]", "droop_calc": "Droop Calc [%]", "aporta": "¿Aporta?"
    }

    tbl = doc.add_table(rows=1 + len(hist_sub), cols=len(cols_disp))
    tbl.style = "Table Grid"
    _add_header_row(tbl, [headers_map.get(c, c) for c in cols_disp])

    for i, (_, fila) in enumerate(hist_sub.iterrows(), start=1):
        row = tbl.rows[i]
        for j, col in enumerate(cols_disp):
            val = fila[col]
            if col == "aporta":
                txt = "✔ Sí" if bool(val) else "✘ No"
                bg  = _COL_OK_BG if bool(val) else _COL_NOK_BG
                _set_cell_bg(row.cells[j], bg)
                _cell_text(row.cells[j], txt, bold=True, size_pt=8,
                           align=WD_ALIGN_PARAGRAPH.CENTER)
            elif isinstance(val, float):
                _cell_text(row.cells[j], f"{val:.2f}", size_pt=8,
                           align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                _cell_text(row.cells[j], str(val), size_pt=8)

    doc.add_paragraph()
    n_hist = len(hist_sub)
    _add_placeholder_copilot(
        doc,
        f"Con base en los {n_hist} eventos mostrados de {ctx.unit_name}, identifica: "
        "1) si el evento actual es típico o atípico, 2) tendencia de mejora o deterioro "
        "en el aporte RPF, 3) recomendaciones específicas para COBEE."
    )


def _doc_diagnostico(doc: Document, ctx: ReporteContext) -> None:
    doc.add_heading("6. Diagnóstico de Cumplimiento Normativo", level=1)

    # Criterios CNDC
    doc.add_heading("6.1 Criterios CNDC Aplicados", level=2)
    criterios = [
        "ΔP% ≥ 1.5 % respecto a P_max para declarar aporte a la RPF",
        "Rango válido de Droop (CDM): 6 % – 12 %",
        "Banda muerta de frecuencia: ± 25 mHz alrededor de 50 Hz",
        "ROCOF calculado por regresión lineal en ventana [0, 3 s] post-falla",
        f"Ventana de evaluación Δt = {ctx.delta_t} s (configurable en el Bloque 3)",
    ]
    for c in criterios:
        p = doc.add_paragraph(c, style="List Bullet")
        p.runs[0].font.size = Pt(9)

    # Veredicto consolidado
    doc.add_heading("6.2 Veredicto Consolidado", level=2)
    fuentes_kpi = [
        ("SCADA COBEE",         ctx.kpi_scada),
        ("EMF CNDC",            ctx.kpi_emf),
        (f"Sim E{ctx.n_evento}.0", ctx.kpi_sim_e0),
        (f"Sim E{ctx.n_evento}.1", ctx.kpi_sim_e1),
    ]
    activos = [(n, k) for n, k in fuentes_kpi if k is not None]
    if activos:
        n_aporta = sum(1 for _, k in activos if k.get("aporta", False))
        tbl = doc.add_table(rows=1 + len(activos), cols=4)
        tbl.style = "Table Grid"
        _add_header_row(tbl, ["Fuente", "ΔP%", "Droop calc", "Veredicto"])
        for i, (nombre, kpi) in enumerate(activos, start=1):
            aporta = kpi.get("aporta", False)
            bg = _COL_OK_BG if aporta else _COL_NOK_BG
            row = tbl.rows[i]
            _cell_text(row.cells[0], nombre, size_pt=9)
            _cell_text(row.cells[1], f"{kpi.get('dp_pct', 0):.2f}%",
                       size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            dc = kpi.get("droop_calc", "—")
            _cell_text(row.cells[2], f"{dc:.1f}%" if isinstance(dc, float) else str(dc),
                       size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_bg(row.cells[3], bg)
            _cell_text(row.cells[3], "APORTA RPF" if aporta else "NO APORTA",
                       bold=True, size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph()
        p_res = doc.add_paragraph(
            f"Resultado: {n_aporta} de {len(activos)} fuentes confirman aporte a la RPF."
        )
        p_res.runs[0].bold = True
        p_res.runs[0].font.size = Pt(10)

    _add_placeholder_copilot(
        doc,
        "Proporciona un diagnóstico técnico detallado (máx. 200 palabras): evalúa si los "
        "resultados de droop calculado son consistentes con el CDM, si el ROCOF refleja "
        "una adecuada inercia, e identifica cualquier inconsistencia normativa."
    )


def _doc_recomendaciones(doc: Document, ctx: ReporteContext) -> None:
    doc.add_heading("7. Recomendaciones", level=1)
    _add_placeholder_copilot(
        doc,
        "Genera 3 a 5 recomendaciones técnicas específicas para COBEE basadas en los "
        f"resultados del análisis de {ctx.unit_name}. Incluye: ajuste de parámetros de "
        "regulación, revisión de datos de entrada, o acciones ante diferencias entre "
        "simulación y mediciones reales. Tono técnico-regulatorio."
    )


def _doc_apendice_metodologia(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Apéndice A — Metodología CNDC RPF", level=1)
    doc.add_paragraph(
        "La evaluación de la Reserva de Potencia de Frecuencia (RPF) sigue la "
        "metodología oficial del Comité Nacional de Despacho de Carga (CNDC) de Bolivia."
    ).runs[0].font.size = Pt(9)

    pasos = [
        ("Paso 1 — Identificación del Evento",
         "Se identifica el instante t₀ de inicio del evento de frecuencia (caída ≥ umbral df/dt) "
         "y se define la ventana de análisis [t₀, t₀+Δt]."),
        ("Paso 2 — Medición de Indicadores",
         "Se registran: f₀ (frecuencia al inicio), f_min (nadir, mínimo de frecuencia "
         "post-falla), P₀ (potencia al inicio) y P_Δt (potencia al tiempo t₀+Δt)."),
        ("Paso 3 — Cálculo del Aporte",
         "ΔP = P_Δt − P₀ [MW]. ΔP% = ΔP / P_max × 100 [%]. "
         "La unidad APORTA si ΔP% ≥ 1.5%."),
        ("Paso 4 — Verificación del Droop",
         "Droop calculado = (Δf'/f_nom) / (ΔP/P_max) × 100 [%], con Δf' aplicando "
         "la banda muerta de ±25 mHz. Se compara con el droop nominal declarado en el CDM."),
        ("ROCOF",
         "La tasa de cambio de frecuencia se calcula por regresión lineal en la ventana "
         "[0, 3 s] post-falla como indicador de la inercia efectiva del sistema."),
    ]
    for titulo, texto in pasos:
        p_t = doc.add_paragraph()
        run_t = p_t.add_run(titulo)
        run_t.bold = True
        run_t.font.size = Pt(9)
        p_b = doc.add_paragraph(texto)
        p_b.runs[0].font.size = Pt(9)
        p_b.paragraph_format.left_indent = Cm(0.5)


# ── Función principal Word ────────────────────────────────────────────────────

def generar_reporte_word(ctx: ReporteContext,
                          incluir_historico: bool = True,
                          incluir_apendice: bool = False) -> bytes:
    """
    Genera el documento Word completo. Retorna bytes del .docx.
    Pasar a st.download_button(data=...) o a subir_reporte_a_sharepoint().
    """
    doc = Document()

    # Márgenes de página
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # Tamaño fuente base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    _doc_portada(doc, ctx)
    _doc_resumen_ejecutivo(doc, ctx)
    doc.add_paragraph()
    _doc_contexto_evento(doc, ctx)
    doc.add_paragraph()
    _doc_resultados_por_fuente(doc, ctx)
    _doc_tabla_comparativa(doc, ctx)
    doc.add_paragraph()
    if incluir_historico:
        _doc_historico(doc, ctx)
        doc.add_paragraph()
    _doc_diagnostico(doc, ctx)
    doc.add_paragraph()
    _doc_recomendaciones(doc, ctx)
    _doc_apendice_metodologia(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Copilot Context Pack ──────────────────────────────────────────────────────

def generar_copilot_context_pack(ctx: ReporteContext) -> str:
    """
    Genera texto estructurado + prompts pre-escritos para copiar/pegar en Copilot.
    """
    lines = []
    lines.append(f"# CONTEXTO RPF — {ctx.semestre} — EVENTO {ctx.n_evento} — {ctx.unit_name}")
    lines.append(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}  |  Sistema: SIN Bolivia  |  Empresa: COBEE")
    lines.append("")
    lines.append("## PARÁMETROS DE LA UNIDAD")
    lines.append(f"- Unidad: {ctx.unit_name}")
    lines.append(f"- P_max: {ctx.p_max:.2f} MW")
    lines.append(f"- Rp nominal (estatismo): {ctx.rp_nominal * 100:.1f}%")
    lines.append(f"- Δt evaluación CNDC: {ctx.delta_t} s")
    lines.append("")

    def _kpi_linea(nombre: str, kpi: Optional[dict]) -> str:
        if kpi is None:
            return f"- {nombre}: No disponible"
        aporta_txt = "✔ APORTA" if kpi.get("aporta", False) else "✘ NO APORTA"
        dp_pct = kpi.get("dp_pct", 0)
        f_min  = kpi.get("f_min", 0)
        droop  = kpi.get("droop_calc", "—")
        rocof  = kpi.get("rocof")
        rocof_txt = f"{rocof:.4f} Hz/s" if rocof is not None and not (isinstance(rocof, float) and np.isnan(rocof)) else "—"
        droop_txt = f"{droop:.1f}%" if isinstance(droop, float) else str(droop)
        return (f"- {nombre}: ΔP%={dp_pct:.2f}% | f_min={f_min:.4f} Hz | "
                f"Droop={droop_txt} | ROCOF={rocof_txt} | {aporta_txt}")

    lines.append("## KPIs POR FUENTE")
    lines.append(_kpi_linea("SCADA COBEE (1 SEG)",         ctx.kpi_scada))
    lines.append(_kpi_linea("EMF CNDC (Digitalización)",   ctx.kpi_emf))
    lines.append(_kpi_linea(f"Simulación E{ctx.n_evento}.0 (CNDC)",   ctx.kpi_sim_e0))
    lines.append(_kpi_linea(f"Simulación E{ctx.n_evento}.1 (COBEE)",  ctx.kpi_sim_e1))
    lines.append("")

    lines.append("## CRITERIOS NORMATIVOS CNDC RPF")
    lines.append("- Aporte si ΔP% ≥ 1.5% respecto a P_max")
    lines.append("- Rango droop válido (CDM): 6% – 12%")
    lines.append("- Banda muerta de frecuencia: ±25 mHz")
    lines.append("- ROCOF en ventana [0, 3 s] post-falla (regresión lineal)")
    lines.append("")

    # Histórico (últimas 5 filas)
    if not ctx.historico.empty:
        hist = ctx.historico.tail(5)
        lines.append(f"## HISTÓRICO DE {ctx.unit_name} (últimos {len(hist)} eventos)")
        for _, row in hist.iterrows():
            sem = row.get("semestre", "—")
            ev  = row.get("evento", "—")
            dp  = row.get("dp_pct", None)
            ap  = "✔" if bool(row.get("aporta", False)) else "✘"
            dp_txt = f"{dp:.2f}%" if dp is not None else "—"
            lines.append(f"  · {sem} Ev{ev}: ΔP%={dp_txt} {ap}")
        lines.append("")
    else:
        lines.append("## HISTÓRICO")
        lines.append("- Sin datos históricos disponibles en la base de datos.")
        lines.append("")

    # Determinar fuente principal para prompts
    kpi_ref = ctx.kpi_scada or ctx.kpi_emf or ctx.kpi_sim_e0 or ctx.kpi_sim_e1
    if kpi_ref:
        fuente_ref = kpi_ref.get("fuente", "fuente principal")
        dp_pct_ref = kpi_ref.get("dp_pct", 0)
        f_min_ref  = kpi_ref.get("f_min", 49.0)
        droop_ref  = kpi_ref.get("droop_calc", "—")
        droop_txt  = f"{droop_ref:.1f}%" if isinstance(droop_ref, float) else str(droop_ref)
        n_hist     = len(ctx.historico)
    else:
        fuente_ref, dp_pct_ref, f_min_ref, droop_txt, n_hist = "datos", 0, 49.0, "—", 0

    lines.append("---")
    lines.append("## PROMPTS SUGERIDOS PARA MICROSOFT COPILOT")
    lines.append("(Copia y pega el prompt que necesites en Teams Copilot Chat o en Word Copilot)")
    lines.append("")

    lines.append("### [PROMPT 1 — Narrativa Técnica]")
    lines.append(
        f'Redacta un párrafo técnico formal (máx. 150 palabras) sobre el comportamiento de '
        f'"{ctx.unit_name}" durante el Evento {ctx.n_evento} ({ctx.semestre}). '
        f'Indica si cumplió con la RPF (ΔP% = {dp_pct_ref:.2f}% según {fuente_ref}), '
        f'el nadir de frecuencia alcanzado ({f_min_ref:.4f} Hz) y el droop calculado '
        f'({droop_txt} vs. nominal {ctx.rp_nominal * 100:.1f}%). '
        f'Tono: informe técnico regulatorio boliviano para presentar al CNDC.'
    )
    lines.append("")

    lines.append("### [PROMPT 2 — Diagnóstico de Cumplimiento]")
    lines.append(
        f'Analiza si "{ctx.unit_name}" cumplió la normativa de Reserva de Potencia de '
        f'Frecuencia (RPF) del CNDC Bolivia en el Evento {ctx.n_evento}. Considera: '
        f'1) Criterio de aporte: ΔP% ≥ 1.5% (resultado: {dp_pct_ref:.2f}%), '
        f'2) Estatismo declarado vs calculado ({ctx.rp_nominal * 100:.1f}% vs {droop_txt}), '
        f'3) Consistencia entre las fuentes disponibles (SCADA, EMF, simulaciones). '
        f'¿Qué factores técnicos podrían explicar las diferencias encontradas?'
    )
    lines.append("")

    lines.append("### [PROMPT 3 — Comparación Histórica]")
    lines.append(
        f'Con base en el historial de "{ctx.unit_name}" ({n_hist} eventos analizados) '
        f'y los resultados del Evento {ctx.n_evento}, identifica: '
        f'1) si el comportamiento del evento actual es típico o atípico para esta unidad, '
        f'2) si existe tendencia de mejora o deterioro en el aporte RPF, '
        f'3) recomendaciones específicas para COBEE respecto a esta unidad generadora.'
    )
    lines.append("")

    lines.append("### [PROMPT 4 — Sección Informe CNDC]")
    lines.append(
        f'Redacta la sección "Análisis de {ctx.unit_name}" para el Informe RPF {ctx.semestre} '
        f'a entregar al CNDC Bolivia. La sección debe incluir: '
        f'a) Identificación de la unidad y el evento, '
        f'b) Tabla de resultados con los KPIs principales, '
        f'c) Veredicto de cumplimiento con justificación técnica, '
        f'd) Observaciones sobre droop y ROCOF. '
        f'Formato: máx. 1 página A4, tono técnico regulatorio formal.'
    )

    return "\n".join(lines)


# ── Subida a SharePoint ───────────────────────────────────────────────────────

def subir_reporte_a_sharepoint(
    docx_bytes:  bytes,
    ev_path:     str,
    n_evento,
    unit_name:   str,
    semestre:    str,
    upload_fn,              # sharepoint_client.upload_file(local_path, sp_folder)
    local_to_sp_fn,         # sharepoint_client._local_to_sp_folder o equivalente
    raiz_local:  str = "",
) -> tuple[bool, str]:
    """
    Guarda el .docx en ev_path/Reportes/ y lo sube a SharePoint.
    Retorna (exito, mensaje).
    """
    fname = f"RPF_{semestre}_Ev{n_evento}_{unit_name.replace(' ', '_')}.docx"
    rep_dir = os.path.join(ev_path, "Reportes")
    os.makedirs(rep_dir, exist_ok=True)
    local_path = os.path.join(rep_dir, fname)

    try:
        with open(local_path, "wb") as f:
            f.write(docx_bytes)
    except Exception as e:
        return False, f"Error al guardar localmente: {e}"

    try:
        sp_folder = local_to_sp_fn(rep_dir) if local_to_sp_fn else None
        if sp_folder is None:
            return False, "No se pudo determinar la carpeta SharePoint."
        upload_fn(local_path, sp_folder)
        return True, f"Subido a SharePoint: {fname}"
    except Exception as e:
        return False, f"Error en la subida a SharePoint: {e}"
